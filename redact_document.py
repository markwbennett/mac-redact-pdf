#!/usr/bin/env python3
"""
Document Redaction Tool with Claude AI Integration
Redacts client information from PDF and DOCX files using Claude CLI to identify sensitive data.
"""

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import os
import sys
import subprocess
import json
import argparse
import tempfile
import re
from pathlib import Path

# Try to import python-docx for DOCX support
try:
    from docx import Document
    from docx.shared import RgbColor
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


def extract_text_from_pdf(pdf_path):
    """Extract all text from a PDF file for analysis."""
    doc = fitz.open(pdf_path)
    text_parts = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if text.strip():
            text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

    doc.close()
    return "\n\n".join(text_parts)


def extract_text_from_docx(docx_path):
    """Extract all text from a DOCX file for analysis."""
    if not DOCX_SUPPORT:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document(docx_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def identify_redactions_with_claude(text, additional_terms=None):
    """
    Use Claude CLI to identify client information that should be redacted.

    Args:
        text: The document text to analyze
        additional_terms: Optional list of additional terms to include

    Returns:
        List of terms to redact
    """
    # Create a prompt for Claude to identify sensitive information
    prompt = """Analyze this legal document and identify ALL client-identifying information that should be redacted for sharing. Include:

1. **Client names** - Full names, first names, last names, nicknames, aliases
2. **Case numbers** - Any format (e.g., 4:23-cr-00123, 2024-CF-001234, CR-2024-0001)
3. **Docket numbers** - Any court docket references
4. **Addresses** - Street addresses, cities (when identifying), zip codes
5. **Phone numbers** - Any format
6. **Email addresses**
7. **Social Security numbers** - Any format (XXX-XX-XXXX or variations)
8. **Dates of birth**
9. **Driver's license numbers**
10. **Account numbers** - Bank, financial, medical record numbers
11. **Family member names** - Spouses, children, relatives mentioned
12. **Employer names** - When clearly associated with the client
13. **Specific identifying details** - Unique identifiers that could identify the client

IMPORTANT:
- Return ONLY a JSON array of strings, one term per array element
- Include variations (e.g., "John Smith", "Smith", "John", "Mr. Smith")
- Be thorough - when in doubt, include it
- Do NOT include generic legal terms, court names, judge names, or attorney names
- Focus on CLIENT identifying information only

Example output format:
["John Smith", "Smith", "4:23-cr-00123", "123 Main Street", "555-123-4567"]

Document text to analyze:
---
"""

    # Truncate text if too long (Claude has context limits)
    max_chars = 50000
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[Document truncated for analysis...]"

    full_prompt = prompt + text + "\n---\n\nReturn ONLY the JSON array of terms to redact:"

    try:
        # Call Claude CLI
        result = subprocess.run(
            ['claude', '-p', full_prompt],
            capture_output=True,
            text=True,
            timeout=120
        )

        if result.returncode != 0:
            print(f"Warning: Claude CLI returned error: {result.stderr}")
            return additional_terms or []

        response = result.stdout.strip()

        # Try to extract JSON array from response
        # Claude might include extra text, so we look for the array
        json_match = re.search(r'\[.*?\]', response, re.DOTALL)
        if json_match:
            terms = json.loads(json_match.group())
            if isinstance(terms, list):
                # Add any additional terms
                if additional_terms:
                    terms.extend(additional_terms)
                # Remove duplicates while preserving order
                seen = set()
                unique_terms = []
                for term in terms:
                    term_lower = term.lower()
                    if term_lower not in seen and term.strip():
                        seen.add(term_lower)
                        unique_terms.append(term.strip())
                return unique_terms

        print(f"Warning: Could not parse Claude response as JSON array")
        print(f"Response: {response[:500]}")
        return additional_terms or []

    except subprocess.TimeoutExpired:
        print("Warning: Claude CLI timed out")
        return additional_terms or []
    except FileNotFoundError:
        print("Error: Claude CLI not found. Please install it: https://claude.ai/download")
        return additional_terms or []
    except json.JSONDecodeError as e:
        print(f"Warning: Could not parse Claude response: {e}")
        return additional_terms or []


def detect_page_type(page):
    """Detect if a page is native digital content or a scanned image."""
    text = page.get_text().strip()
    images = page.get_images()

    page_rect = page.rect
    page_area = page_rect.width * page_rect.height

    image_area = 0
    for img_info in images:
        try:
            img_rect = page.get_image_bbox(img_info[7])
            if img_rect:
                image_area += img_rect.width * img_rect.height
        except:
            pass

    image_coverage = image_area / page_area if page_area > 0 else 0

    if len(text) > 100 and image_coverage < 0.5:
        return 'native'
    elif len(text) < 50 and image_coverage > 0.8:
        return 'scanned'
    elif image_coverage > 0.8:
        return 'scanned'
    elif len(text) > 50:
        return 'native'
    else:
        return 'scanned'


def strip_text_layer(doc, page_numbers):
    """Remove existing text from specified pages of the PDF."""
    if not page_numbers:
        return

    print(f"  Removing text layer from {len(page_numbers)} scanned page(s)...")

    for page_num in page_numbers:
        page = doc[page_num]
        text_dict = page.get_text("dict")
        blocks = text_dict.get("blocks", [])

        for block in blocks:
            if block.get("type") == 0:
                bbox = fitz.Rect(block["bbox"])
                page.add_redact_annot(bbox)

        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)


def ocr_page(page, dpi=300):
    """Perform OCR on a PDF page and return text with bounding boxes."""
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat)

    img_data = pix.tobytes("png")
    img = Image.open(io.BytesIO(img_data))

    ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

    words_with_boxes = []
    n_boxes = len(ocr_data['text'])

    for i in range(n_boxes):
        text = ocr_data['text'][i].strip()
        if text:
            x0 = ocr_data['left'][i] * 72 / dpi
            y0 = ocr_data['top'][i] * 72 / dpi
            x1 = (ocr_data['left'][i] + ocr_data['width'][i]) * 72 / dpi
            y1 = (ocr_data['top'][i] + ocr_data['height'][i]) * 72 / dpi

            bbox = fitz.Rect(x0, y0, x1, y1)
            words_with_boxes.append((text, bbox))

    return words_with_boxes


def add_ocr_text_layer(doc, page_numbers):
    """Add OCR'd text as an invisible text layer to specified pages."""
    if not page_numbers:
        return {}

    print(f"  Performing OCR on {len(page_numbers)} scanned page(s)...")

    page_ocr_data = {}

    for page_num in page_numbers:
        page = doc[page_num]
        print(f"    OCR'ing page {page_num + 1}/{len(doc)}...")

        words_with_boxes = ocr_page(page)
        page_ocr_data[page_num] = words_with_boxes

        for text, bbox in words_with_boxes:
            page.insert_text(
                bbox.tl,
                text,
                fontsize=bbox.height * 0.8,
                color=(0, 0, 0),
                render_mode=3
            )

    return page_ocr_data


def redact_terms_in_pdf(doc, page_ocr_data, terms):
    """Search for terms in PDF and redact them with black boxes."""
    print(f"  Searching for {len(terms)} term(s) to redact...")

    total_redactions = 0

    for page_num in range(len(doc)):
        page = doc[page_num]
        ocr_words = page_ocr_data.get(page_num, [])

        # Search using PyMuPDF's built-in search for native pages
        for term in terms:
            # Search in native text
            text_instances = page.search_for(term)
            for inst in text_instances:
                page.add_redact_annot(inst, fill=(0, 0, 0))
                total_redactions += 1
                print(f"    Page {page_num + 1}: Redacting '{term}'")

            # Also search in OCR'd text
            term_lower = term.lower()
            for text, bbox in ocr_words:
                if term_lower in text.lower():
                    page.add_redact_annot(bbox, fill=(0, 0, 0))
                    total_redactions += 1

        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    return total_redactions


def redact_pdf(pdf_path, terms, output_path=None):
    """Process and redact a PDF file."""
    print(f"\nProcessing PDF: {pdf_path}")
    doc = fitz.open(pdf_path)

    # Analyze pages
    print("  Analyzing PDF structure...")
    scanned_pages = []
    native_pages = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_type = detect_page_type(page)

        if page_type == 'scanned':
            scanned_pages.append(page_num)
        else:
            native_pages.append(page_num)

    print(f"    Found {len(native_pages)} native page(s), {len(scanned_pages)} scanned page(s)")

    # Process scanned pages
    strip_text_layer(doc, scanned_pages)
    page_ocr_data = add_ocr_text_layer(doc, scanned_pages)

    # Redact terms
    total_redactions = redact_terms_in_pdf(doc, page_ocr_data, terms)

    # Generate output path
    if not output_path:
        base_name, ext = os.path.splitext(pdf_path)
        output_path = f"{base_name}_redacted{ext}"

    # Save
    print(f"\n  Applied {total_redactions} redaction(s)")
    print(f"  Saving to: {output_path}")
    doc.save(output_path, garbage=4, deflate=True)
    doc.close()

    return output_path


def redact_docx(docx_path, terms, output_path=None):
    """Process and redact a DOCX file."""
    if not DOCX_SUPPORT:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    print(f"\nProcessing DOCX: {docx_path}")
    doc = Document(docx_path)

    total_redactions = 0
    redact_marker = "[REDACTED]"

    # Process paragraphs
    for para in doc.paragraphs:
        original_text = para.text
        new_text = original_text

        for term in terms:
            if term.lower() in new_text.lower():
                # Case-insensitive replacement
                pattern = re.compile(re.escape(term), re.IGNORECASE)
                matches = pattern.findall(new_text)
                total_redactions += len(matches)
                new_text = pattern.sub(redact_marker, new_text)

        if new_text != original_text:
            # Clear and rewrite paragraph
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                new_text = original_text

                for term in terms:
                    if term.lower() in new_text.lower():
                        pattern = re.compile(re.escape(term), re.IGNORECASE)
                        matches = pattern.findall(new_text)
                        total_redactions += len(matches)
                        new_text = pattern.sub(redact_marker, new_text)

                if new_text != original_text:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = new_text
                        else:
                            para.add_run(new_text)
                        break

    # Generate output path
    if not output_path:
        base_name, ext = os.path.splitext(docx_path)
        output_path = f"{base_name}_redacted{ext}"

    print(f"\n  Applied {total_redactions} redaction(s)")
    print(f"  Saving to: {output_path}")
    doc.save(output_path)

    return output_path


def process_document(file_path, terms=None, output_path=None, use_claude=True, additional_terms=None):
    """
    Process a document (PDF or DOCX) and redact sensitive information.

    Args:
        file_path: Path to the document
        terms: List of terms to redact (if None, uses Claude to identify)
        output_path: Custom output path (optional)
        use_claude: Whether to use Claude CLI to identify terms
        additional_terms: Additional terms to always include

    Returns:
        Path to the redacted document
    """
    file_path = os.path.expanduser(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    if ext not in ['.pdf', '.docx']:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx")

    # Extract text for Claude analysis
    if use_claude and not terms:
        print("\nExtracting text for analysis...")
        if ext == '.pdf':
            text = extract_text_from_pdf(file_path)
        else:
            text = extract_text_from_docx(file_path)

        print("Asking Claude to identify client information...")
        terms = identify_redactions_with_claude(text, additional_terms)

        if not terms:
            print("\nNo terms identified for redaction.")
            return None

        print(f"\nIdentified {len(terms)} term(s) to redact:")
        for term in terms[:20]:  # Show first 20
            print(f"  - {term}")
        if len(terms) > 20:
            print(f"  ... and {len(terms) - 20} more")
    elif additional_terms:
        terms = list(terms or []) + list(additional_terms)

    if not terms:
        print("No terms to redact.")
        return None

    # Process the document
    if ext == '.pdf':
        return redact_pdf(file_path, terms, output_path)
    else:
        return redact_docx(file_path, terms, output_path)


def main():
    parser = argparse.ArgumentParser(
        description='Redact client information from PDF and DOCX files using Claude AI',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Auto-detect client info with Claude and redact
  %(prog)s document.pdf

  # Redact specific terms only
  %(prog)s document.pdf --terms "John Smith" "123-45-6789"

  # Add extra terms to Claude's detection
  %(prog)s document.pdf --add-terms "Case XYZ"

  # Specify output file
  %(prog)s document.pdf -o redacted_document.pdf
        """
    )

    parser.add_argument('file', help='Path to PDF or DOCX file to redact')
    parser.add_argument('-o', '--output', help='Output file path (default: adds _redacted suffix)')
    parser.add_argument('--terms', nargs='+', help='Specific terms to redact (skips Claude analysis)')
    parser.add_argument('--add-terms', nargs='+', help='Additional terms to add to Claude\'s detection')
    parser.add_argument('--no-claude', action='store_true', help='Skip Claude analysis (requires --terms)')

    args = parser.parse_args()

    if args.no_claude and not args.terms:
        parser.error("--no-claude requires --terms to specify what to redact")

    print("=" * 60)
    print("Document Redaction Tool with Claude AI")
    print("=" * 60)

    try:
        output = process_document(
            args.file,
            terms=args.terms,
            output_path=args.output,
            use_claude=not args.no_claude,
            additional_terms=args.add_terms
        )

        if output:
            print(f"\n{'=' * 60}")
            print("Redaction complete!")
            print(f"Redacted file: {output}")
            print("=" * 60)

    except Exception as e:
        print(f"\nError: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
